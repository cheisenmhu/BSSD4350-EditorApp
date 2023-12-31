# Programmers: Anita Martin, Chris Heise, Dion Boyer, and Felix Jaramillo
# Course: BSSD 4350 - Agile Methodologies
# Instructor: Jonathan Lee
# Program: Inclusivity Editor App
# File: app.py

import gradio as gr
from difflib import Differ
from fpdf import *
from docx import *
import pyperclip
import together


together.api_key = ""

users_text = ""

EXAMPLE_TEXT = ""

CORRECTED_TEXT = ""

textInput = "There once was a farmer named..."

selected = ""

DIFFERENCES = []

# Global Components (accessed by multiple tabs/pages)
original_text = gr.Textbox(
    label="Your Text",
    info="Your original text.",
)

previewText = gr.Textbox(label="Output Textbox", value=textInput)

# Isn't currently working. Seems to need to be called with a button click like other componenets/functions
# Source: https://github.com/gradio-app/gradio/issues/2412

# Chris
def change_page(page_number):
    """Changes the page to the page number passed in."""
    return gr.Tabs.update(selected=page_number)

# Dion
def update_preview(text):
    gr.Textbox(label="Output Textbox", value=textInput)
    return text

# Dion
def download(output, type):
    match type:
        case "PDF":
            """Download text as a PDF."""
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            pdf.multi_cell(0, 10, str(output))
            pdf_file = "history_download.pdf"
            pdf.output(pdf_file)
            return pdf_file
        case "DOCX":
            doc = Document()

            doc.add_heading('History Download', 0)
            doc.add_paragraph(output)
            doc_output = "history_download.docx"
            doc.save(doc_output)

            return doc_output
        case "TXT":
            text_file = 'history_download.txt'
            with open('history_download.txt', 'w') as txt_file:
                txt_file.write(output)
                txt_file.close()
            return text_file
        case _:
            """Download text as a PDF."""
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            pdf.multi_cell(0, 10, str(output))
            pdf_file = type + "history_download.pdf"
            pdf.output(pdf_file)
            return pdf_file


# Dion
def copy_text(output):
    """Copy text to the clipboard."""
    pyperclip.copy(output)
    text = pyperclip.paste()

# Chris
def load_text(temp_file):
    """Load text from a temporary file."""
    content = ""
    with open(temp_file.name, "r", encoding="utf-8") as f:
        content = f.read()
    return content

# Anita
def submit_text(text):
    global textInput, users_text
    users_text = text
    textInput += users_text
    change_page(2)
    return users_text

# Chris
def diff_texts(text1, text2):
    """Find the differences between two texts."""
    d = Differ()
    return [
        (token[2:], token[0] if token[0] != " " else None)
        for token in d.compare(text1, text2)
    ]

# Dion
def dropdown_callback(value):
    return value

# Felix
def prompts(choice):
    global selected
    if choice == "Search for Grammar Errors":
        selected = "Do not rewrite, just fix the grammatical errors in the following sentences:"
        return selected
    elif choice == "Professional Correspondence":
        selected = "Rewrite the following as a Professional Correspondence"
        return selected
    elif choice == "Personal Correspondence":
        selected = "Rewrite the following as a Personal Correspondence"
        return selected
    elif choice == "Educational Paper":
        selected = "Rewrite the following as an Educational Paper"
        return selected
    elif choice == "Technical Instructions":
        selected = "Rewrite the following to Technical Instructions"
        return selected

# Anita
def call_llm(prompt_text):
    llm = together.Complete.create(
        prompt=selected + " " + prompt_text,
        model="togethercomputer/llama-2-7b-chat",
        max_tokens=256,
        temperature=0.8,
        top_k=60,
        top_p=0.6,
        repetition_penalty=1.1,
        stop=['<human>']
    )
    answer = (llm['output']['choices'][0]['text']).strip().split("Answer:\n")[0]
    return answer

# Chris, Anita, Felix, Dion
with gr.Blocks() as incluesive:
    gr.Markdown("# INCLUeSIVE")
    with gr.Tabs() as pages:
        """FIRST PAGE"""
        # Felix
        with gr.TabItem("Welcome", id=0) as first_page:
            with gr.Tab("Writing Preferences"):
                gr.Markdown(
                    "Welcome to Incluesive an app that will help correcct your writings to be more incluesive of everyone. "
                    "To use Incluesive Pick a writing purpose then enter your text into the text box and submit. After you submit the changes to your text will be shown.")
                choice = gr.Radio(
                    ["Search for Grammar Errors", "Professional Correspondence", "Personal Correspondence",
                     "Educational Paper", "Technical Instructions"], label="Writing purpose")
                choice.change(fn=prompts, inputs=choice, outputs=None)
        """END FIRST PAGE"""

        """SECOND PAGE"""
        # Chris and Anita
        with gr.TabItem("Input", id=1) as second_page:
            with gr.Tab("Type/Paste"):
                text_input = gr.Textbox(
                    label="Type or paste your text here.",
                    info="Your Original Text.",
                    value=EXAMPLE_TEXT,
                )
                submit_button = gr.Button("Submit")
                corrected_text = gr.Textbox(
                    label="Corrected Text",
                    info="Our suggested corrected text",
                    value=CORRECTED_TEXT,
                    visible=False,
                )
                submit_button.click(submit_text, inputs=[text_input], outputs=original_text)
                clear_button = gr.ClearButton(text_input)
            with gr.Tab("Upload"):
                file_input = gr.File(
                    file_types=["text"],
                )
                with gr.Row():
                    upload_button = gr.Button("Upload")
                loaded_text = gr.Textbox(
                    label="Your Text",
                    info="The text you uploaded.",
                )
                with gr.Row():
                    clear_button = gr.ClearButton(loaded_text)
                    submit_button = gr.Button("Submit")
                    upload_button.click(load_text, inputs=[file_input], outputs=[loaded_text])
                    submit_button.click(submit_text, inputs=[loaded_text], outputs=original_text)
        """END SECOND PAGE"""

        """THIRD PAGE"""
        # Anita
        with gr.TabItem("Results", id=2) as third_page:
            input_text = original_text.render()
            output_text = gr.Textbox(label="Results from LLM")
            output_text.change(update_preview, output_text, previewText)
            with gr.Row():
                submit_button = gr.Button("Make Request")
                clear_button = gr.ClearButton(original_text)

            corrections = gr.HighlightedText(
                label="Corrections",
                combine_adjacent=True,
                show_legend=True,
                color_map={"+": "green", "-": "red"},
            )
            with gr.Row():
                highlight_button = gr.Button("Highlight Differences")
                accept_paragraph_button = gr.Button("Ignore")
                submit_paragraph_button = gr.Button("Submit")

            submit_button.click(fn=call_llm, inputs=input_text, outputs=output_text)
            highlight_button.click(diff_texts, inputs=[input_text, output_text], outputs=[corrections])

        """END THIRD PAGE"""

        """FOURTH PAGE"""
        # Dion
        with gr.TabItem("Save", id=3) as fourth_page:
            with gr.Accordion(label="Account"):
                preferences = gr.Button(value="Preferences")
                signout = gr.Button(value="Sign Out")
            with gr.Row():
                previewText.render()
                file = gr.File()
            with gr.Row():
                dropdown_type = gr.Dropdown(
                    ["DOCX", "PDF", "TXT"], label="File Type", info="Select your file type."
                )
                download_btn = gr.Button(value="Download", scale=0)
                copy_btn = gr.Button(value="Copy", scale=0)
                done_btn = gr.Button(value="Done", scale=0)
                download_btn.click(fn=download, inputs=[previewText, dropdown_type], outputs=file, api_name="Download")
                copy_btn.click(fn=copy_text, inputs=previewText, api_name="Copy")
        """END FOURTH PAGE"""

incluesive.launch()
